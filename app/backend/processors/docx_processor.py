"""
文档任务处理器

负责处理Word文档润色任务
"""
from __future__ import annotations

from typing import Optional

import os
from backend import paths
import docx
from concurrent.futures import ThreadPoolExecutor, as_completed
from backend.processors.base_processor import BaseTaskProcessor
from backend.utils.docx_service import is_paragraph_needs_polishing, replace_paragraph_text, check_stop_signal
from backend.utils.helpers import split_text_into_chunks
from backend.config import WorkerConfig, RedisKeyManager
from backend.extensions import db
from backend.utils.logging_config import get_logger

logger = get_logger(__name__)


class DocxTaskProcessor(BaseTaskProcessor):
    """文档任务处理器"""

    def __init__(self, task: object, redis_client: object) -> None:
        super().__init__(task, redis_client)
        self.doc: Optional[docx.Document] = None

    def process(self) -> None:
        """处理文档任务"""
        done_indices_key = RedisKeyManager.docx_done_key(self.task_id)

        # 从Redis恢复已完成的段落索引
        raw_done = self.redis_client.smembers(done_indices_key)
        done_indices = {int(x) for x in raw_done} if raw_done else set()

        # 加载文档（如果有中间结果，加载中间结果）
        file_to_load = self.task.result_file_path if (
            self.task.result_file_path and len(done_indices) > 0
        ) else self.task.file_path

        self.doc = docx.Document(file_to_load)
        logger.info(f"任务 {self.task_id} 加载文档: {file_to_load}")

        # 提取需要处理的段落
        paras_to_process = self._extract_paragraphs(done_indices)
        total_paras = len(paras_to_process)

        if total_paras == 0:
            logger.info(f"任务 {self.task_id} 没有需要处理的段落")
            self._save_document()
            return

        logger.info(f"任务 {self.task_id} 需要处理 {total_paras} 个段落")

        self.progress_publisher.publish_block(
            f"📄 解析成功！共提取 {total_paras} 个核心段落待处理。\n"
            f"🚀 【多线程并发】引擎已激活，正在处理...\n\n"
        )

        # 并发处理段落
        self._process_paragraphs_concurrent(paras_to_process, total_paras, done_indices_key)

        # 保存文档
        self._save_document()

        # 清理Redis缓存
        self.redis_client.delete(done_indices_key)
        self.redis_client.delete(RedisKeyManager.docx_progress_key(self.task_id))

        self.progress_publisher.publish_block("\n\n🎉 全局高并发重排完毕，正在生成下载链接...\n")

    def _extract_paragraphs(self, done_indices: set[int]) -> list[tuple[int, str]]:
        """
        提取需要处理的段落

        Args:
            done_indices: 已完成的段落索引集合

        Returns:
            list: [(索引, 文本), ...]
        """
        paras_to_process = []
        stop_processing = False

        for i, para in enumerate(self.doc.paragraphs):
            if stop_processing:
                break

            # 检查停止信号
            if check_stop_signal(para.text):
                stop_processing = True
                break

            # 跳过已完成的段落
            if i in done_indices:
                continue

            # 检查是否需要润色
            if is_paragraph_needs_polishing(para, self.task.mode):
                paras_to_process.append((i, para.text))

        return paras_to_process

    def _process_paragraphs_concurrent(self, paras_to_process: list[tuple[int, str]], total_paras: int, done_indices_key: str) -> None:
        """
        并发处理段落

        Args:
            paras_to_process: 待处理段落列表
            total_paras: 总段落数
            done_indices_key: Redis完成索引键
        """
        results_dict = {}
        completed_count = 0

        with ThreadPoolExecutor(max_workers=WorkerConfig.MAX_WORKERS) as executor:
            # 提交所有任务
            future_to_para = {
                executor.submit(self._process_single_paragraph, idx, text): (idx, text)
                for idx, text in paras_to_process
            }

            # 处理完成的任务
            for future in as_completed(future_to_para):
                # 检查取消信号
                if self.check_cancellation():
                    # 保存中间结果
                    self._apply_results(results_dict)
                    self._save_temp_document()
                    return

                idx, original_text = future_to_para[future]

                try:
                    polished_text = future.result()

                    # 如果返回结果太短，使用原文
                    if not polished_text or len(polished_text.strip()) < 5:
                        polished_text = original_text

                    results_dict[idx] = polished_text
                    completed_count += 1

                    # 记录到Redis
                    self.redis_client.sadd(done_indices_key, idx)

                    # 推送进度
                    progress = int((completed_count / total_paras) * 100)
                    self.progress_publisher.publish_progress(progress)
                    self.progress_publisher.publish_block(
                        f"⚡ [并发加速] 第 {completed_count}/{total_paras} 个段落已完成覆盖...\n"
                    )

                except Exception as e:
                    logger.error(f"处理段落 {idx} 失败: {str(e)}")
                    # 失败时使用原文
                    results_dict[idx] = original_text

        # 应用所有结果
        self._apply_results(results_dict)

    def _process_single_paragraph(self, para_idx: int, text_content: str) -> str:
        """
        处理单个段落

        Args:
            para_idx: 段落索引
            text_content: 段落文本

        Returns:
            str: 润色后的文本
        """
        logger.debug(f"处理段落 {para_idx}")

        # 如果段落太长，先切片
        if len(text_content) > WorkerConfig.get_chunk_size():
            sub_chunks = split_text_into_chunks(text_content, max_chars=WorkerConfig.get_chunk_size())
            polished_sub = []

            for sc in sub_chunks:
                polished_sub.append(
                    self.ai_service.generate_sync(
                        text=sc,
                        mode=self.task.mode,
                        strategy=self.task.strategy
                    )
                )

            return "".join(polished_sub)
        else:
            return self.ai_service.generate_sync(
                text=text_content,
                mode=self.task.mode,
                strategy=self.task.strategy
            )

    def _apply_results(self, results_dict: dict[int, str]) -> None:
        """
        应用润色结果到文档

        Args:
            results_dict: {段落索引: 润色文本}
        """
        for idx, new_text in results_dict.items():
            replace_paragraph_text(self.doc.paragraphs[idx], new_text)

        logger.info(f"应用了 {len(results_dict)} 个段落的润色结果")

    def _save_document(self) -> None:
        """保存最终文档"""
        output_path = os.path.join(paths.get_output_dir(), f"polished_{self.task_id}.docx")
        self.doc.save(output_path)
        self.task.result_file_path = output_path
        db.session.commit()

        logger.info(f"任务 {self.task_id} 文档已保存: {output_path}")

        self.progress_publisher.publish_download(f"/api/tasks/download/{self.task_id}")

    def _save_temp_document(self) -> None:
        """保存临时文档（任务取消时）"""
        temp_path = os.path.join(paths.get_output_dir(), f"polished_temp_{self.task_id}.docx")
        self.doc.save(temp_path)
        self.task.result_file_path = temp_path
        db.session.commit()

        logger.info(f"任务 {self.task_id} 临时文档已保存: {temp_path}")

    def handle_failure(self, exception: Exception) -> None:
        """处理失败情况"""
        # 保存临时文档
        if self.doc is not None:
            try:
                self._save_temp_document()
                logger.info(f"任务 {self.task_id} 失败，已保存临时文档")
            except Exception as e:
                logger.error(f"保存临时文档失败: {str(e)}")

    def cleanup(self) -> None:
        """清理资源"""
        super().cleanup()
        # 清理进度缓存
        done_indices_key = RedisKeyManager.docx_done_key(self.task_id)
        self.redis_client.delete(done_indices_key)
        self.redis_client.delete(RedisKeyManager.docx_progress_key(self.task_id))
