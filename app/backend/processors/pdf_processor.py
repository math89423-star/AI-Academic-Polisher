"""
PDF文档任务处理器

提取PDF文本，润色后输出为docx
"""
from __future__ import annotations

from typing import Optional

import os
import re
import fitz
import docx
from concurrent.futures import ThreadPoolExecutor, as_completed
from backend.processors.base_processor import BaseTaskProcessor
from backend.utils.docx_service import check_stop_signal
from backend.utils.helpers import split_text_into_chunks
from backend.config import WorkerConfig, RedisKeyManager
from backend.extensions import db
from backend.utils.logging_config import get_logger

logger = get_logger(__name__)

FORMULA_PATTERN = re.compile(
    r'[∀-⋿⨀-⫿⟀-⟯←-⇿]'
    r'|\\(?:frac|sum|int|lim|sqrt|alpha|beta|gamma|delta|theta|lambda|sigma|omega)'
    r'|\$[^$]+\$'
)


def _is_pdf_paragraph_needs_polishing(text: str, mode: str = 'zh') -> bool:
    text = text.strip()
    if len(text) < 30:
        return False
    if FORMULA_PATTERN.search(text):
        return False
    if re.match(r'^(图|表|Figure|Table|Fig\.)\s*\d', text):
        return False
    if mode == 'zh':
        en_chars = len(re.findall(r'[a-zA-Z]', text))
        if len(text) > 0 and (en_chars / len(text)) > 0.8:
            return False
    elif mode == 'en':
        zh_chars = len(re.findall(r'[一-龥]', text))
        if len(text) > 0 and (zh_chars / len(text)) > 0.4:
            return False
    return True


class PdfTaskProcessor(BaseTaskProcessor):

    def __init__(self, task: object, redis_client: object) -> None:
        super().__init__(task, redis_client)
        self.paragraphs: list[str] = []


    def process(self) -> None:
        done_indices_key = RedisKeyManager.docx_done_key(self.task_id)
        raw_done = self.redis_client.smembers(done_indices_key)
        done_indices = {int(x) for x in raw_done} if raw_done else set()

        self._extract_text_from_pdf()
        paras_to_process = [(i, p) for i, p in enumerate(self.paragraphs) if i not in done_indices and _is_pdf_paragraph_needs_polishing(p, self.task.mode)]
        total_paras = len(paras_to_process)

        if total_paras == 0:
            logger.info(f"任务 {self.task_id} 没有需要处理的段落")
            self._save_as_docx({})
            return

        logger.info(f"任务 {self.task_id} 需要处理 {total_paras} 个段落")
        self.progress_publisher.publish_block(f"📄 PDF解析成功！共提取 {total_paras} 个核心段落待处理。\n🚀 【多线程并发】引擎已激活，正在处理...\n\n")

        results_dict = self._process_paragraphs_concurrent(paras_to_process, total_paras, done_indices_key)
        self._save_as_docx(results_dict)

        self.redis_client.delete(done_indices_key)
        self.progress_publisher.publish_block("\n\n🎉 全局高并发重排完毕，正在生成下载链接...\n")


    def _extract_text_from_pdf(self) -> None:
        pdf_doc = fitz.open(self.task.file_path)
        raw_paragraphs = []
        stop = False
        for page in pdf_doc:
            if stop:
                break
            blocks = page.get_text("blocks")
            for b in blocks:
                if b[6] != 0:  # image block
                    continue
                text = b[4].strip()
                if not text:
                    continue
                if check_stop_signal(text):
                    stop = True
                    break
                raw_paragraphs.append(text)
        pdf_doc.close()
        self.paragraphs = raw_paragraphs
        logger.info(f"任务 {self.task_id} 从PDF提取 {len(self.paragraphs)} 个文本块")


    def _process_paragraphs_concurrent(self, paras_to_process: list[tuple[int, str]], total_paras: int, done_indices_key: str) -> dict[int, str]:
        results_dict = {}
        completed_count = 0

        with ThreadPoolExecutor(max_workers=WorkerConfig.MAX_WORKERS) as executor:
            future_to_para = {
                executor.submit(self._process_single_paragraph, idx, text): (idx, text)
                for idx, text in paras_to_process
            }
            for future in as_completed(future_to_para):
                if self.check_cancellation():
                    return results_dict

                idx, original_text = future_to_para[future]
                try:
                    polished_text = future.result()
                    if not polished_text or len(polished_text.strip()) < 5:
                        polished_text = original_text
                    results_dict[idx] = polished_text
                    completed_count += 1
                    self.redis_client.sadd(done_indices_key, idx)
                    progress = int((completed_count / total_paras) * 100)
                    self.progress_publisher.publish_progress(progress)
                    self.progress_publisher.publish_block(f"⚡ [并发加速] 第 {completed_count}/{total_paras} 个段落已完成覆盖...\n")
                except Exception as e:
                    logger.error(f"处理段落 {idx} 失败: {str(e)}")
                    results_dict[idx] = original_text

        return results_dict


    def _process_single_paragraph(self, para_idx: int, text_content: str) -> str:
        if len(text_content) > WorkerConfig.TEXT_CHUNK_SIZE:
            sub_chunks = split_text_into_chunks(text_content, max_chars=WorkerConfig.TEXT_CHUNK_SIZE)
            polished_sub = []
            for sc in sub_chunks:
                polished_sub.append(self.ai_service.generate_sync(text=sc, mode=self.task.mode, strategy=self.task.strategy))
            return "".join(polished_sub)
        else:
            return self.ai_service.generate_sync(text=text_content, mode=self.task.mode, strategy=self.task.strategy)

    def _save_as_docx(self, results_dict: dict[int, str]) -> None:
        doc = docx.Document()
        for i, para_text in enumerate(self.paragraphs):
            text = results_dict.get(i, para_text)
            doc.add_paragraph(text)

        output_path = os.path.join('outputs', f"polished_{self.task_id}.docx")
        doc.save(output_path)
        self.task.result_file_path = output_path
        db.session.commit()
        logger.info(f"任务 {self.task_id} PDF润色结果已保存: {output_path}")
        self.progress_publisher.publish_download(f"/api/tasks/download/{self.task_id}")

    def handle_failure(self, exception: Exception) -> None:
        logger.error(f"任务 {self.task_id} PDF处理失败: {str(exception)}")

    def cleanup(self) -> None:
        super().cleanup()
        done_indices_key = RedisKeyManager.docx_done_key(self.task_id)
        self.redis_client.delete(done_indices_key)
