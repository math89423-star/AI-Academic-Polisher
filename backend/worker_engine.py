import json
import os
import docx
from concurrent.futures import ThreadPoolExecutor, as_completed
from backend.extensions import db, redis_client
from backend import create_app
from backend.model import Task, User, ApiConfig
from backend.utils.ai_service import pure_ai_generator, sync_ai_request
from backend.utils.docx_service import is_paragraph_needs_polishing, replace_paragraph_text, check_stop_signal
from backend.utils.helpers import split_text_into_chunks

app = create_app()

def process_task(task_id):
    with app.app_context():
        task = Task.query.get(task_id)
        if not task: return
        user = User.query.get(task.user_id)
        channel_name = f"stream:task:{task_id}"
        
        print(f"\n=======================================")
        print(f"🚀 [Worker] 开始处理任务 ID: {task_id}")
        print(f"=======================================\n")
        
        api_key, base_url, model_name = None, None, 'gpt-3.5-turbo'
        if user and user.api_config:
            api_key, base_url, model_name = user.api_config.api_key, user.api_config.base_url, user.api_config.model_name
        else:
            default_config = ApiConfig.query.first()
            if default_config:
                api_key, base_url, model_name = default_config.api_key, default_config.base_url, default_config.model_name
        
        if not api_key or not base_url:
            _publish_error(channel_name, "系统未配置可用的 API 线路。")
            task.status = 'failed'
            db.session.commit()
            return

        task.status = 'processing'
        db.session.commit()
        
        full_text = ""
        doc = None
        total_chunks = 0 

        try:
            # ==========================================
            # 分支 1：处理纯文本任务
            # ==========================================
            if getattr(task, 'task_type', 'text') == 'text':
                chunks = split_text_into_chunks(task.original_text, max_chars=600)
                total_chunks = len(chunks)
                
                if total_chunks <= 1:
                    # 🟢 新增：由于 Two-Pass 需要后台思考，给用户发送安抚提示！
                    _publish_message(channel_name, "block", "🧠 系统已启用【二次深度提纯】引擎...\n⏳ AI 正在后台默默进行深度思考与草稿构建（此阶段无输出，约需 10~30 秒，请耐心等待，切勿刷新页面）...\n\n")
                    
                    full_text = task.polished_text or ""
                    generator = pure_ai_generator(
                        text=task.original_text, api_key=api_key, base_url=base_url,
                        model_name=model_name, mode=task.mode, history_text=task.polished_text
                    )
                    for chunk in generator:
                        cancel_key = f"cancel:task:{task_id}"
                        if redis_client.exists(cancel_key):
                            task.status = 'cancelled'
                            task.polished_text = full_text
                            db.session.commit()
                            redis_client.delete(cancel_key)
                            return
                        full_text += chunk
                        _publish_message(channel_name, "stream", chunk)
                    
                    task.polished_text = full_text
                
                else:
                    # 2B. 长文本并发处理
                    progress_key = f"text_progress:task:{task_id}"
                    raw_done = redis_client.hgetall(progress_key)
                    done_dict = {int(k): v for k, v in raw_done.items()} if raw_done else {}
                    chunks_to_process = [(i, txt) for i, txt in enumerate(chunks) if i not in done_dict]
                    
                    if chunks_to_process:
                        _publish_message(channel_name, "block", f"📝 长文本切片成功！共 {total_chunks} 个片段。\n🚀 【多线程并发+深度提纯】引擎已激活（首批片段需要静默思考，请稍候）...\n\n")
                        
                        def process_single_chunk(chunk_idx, text_content):
                            res = sync_ai_request(text=text_content, api_key=api_key, base_url=base_url, model_name=model_name, mode=task.mode)
                            return chunk_idx, res

                        completed_count = len(done_dict)
                        is_cancelled = False
                        
                        with ThreadPoolExecutor(max_workers=15) as executor:
                            future_to_chunk = {executor.submit(process_single_chunk, p[0], p[1]): p for p in chunks_to_process}
                            for future in as_completed(future_to_chunk):
                                cancel_key = f"cancel:task:{task_id}"
                                if redis_client.exists(cancel_key):
                                    is_cancelled = True
                                    task.status = 'cancelled'
                                    db.session.commit()
                                    redis_client.delete(cancel_key)
                                    break
                                
                                idx = future_to_chunk[future][0]
                                original_chunk = future_to_chunk[future][1]
                                
                                try:
                                    _, polished_chunk = future.result()
                                    if not polished_chunk or len(polished_chunk.strip()) < 5: polished_chunk = original_chunk
                                    done_dict[idx] = polished_chunk
                                    redis_client.hset(progress_key, str(idx), polished_chunk)
                                    completed_count += 1
                                    
                                    _publish_message(channel_name, "progress", int((completed_count / total_chunks) * 100))
                                    _publish_message(channel_name, "block", f"⚡ [并发加速] 第 {completed_count}/{total_chunks} 个片段完成...\n")
                                except Exception:
                                    done_dict[idx] = original_chunk 
                                    completed_count += 1
                        
                        if is_cancelled: return
                            
                    final_text_parts = [done_dict.get(i, chunks[i]) for i in range(total_chunks)]
                    final_text = "\n\n".join(final_text_parts) 
                    task.polished_text = final_text
                    _publish_message(channel_name, "stream", final_text)
                    redis_client.delete(progress_key)

            # ==========================================
            # 分支 2：DOCX 文档任务
            # ==========================================
            elif getattr(task, 'task_type', '') == 'docx':
                done_indices_key = f"docx_done_indices:task:{task_id}"
                raw_done = redis_client.smembers(done_indices_key)
                done_indices = {int(x) for x in raw_done} if raw_done else set()
                
                file_to_load = task.result_file_path if (task.result_file_path and len(done_indices) > 0) else task.file_path
                doc = docx.Document(file_to_load)
                
                paras_to_process = []
                stop_processing = False
                for i, para in enumerate(doc.paragraphs):
                    if stop_processing: break
                    if check_stop_signal(para.text): 
                        stop_processing = True
                        break
                    if i in done_indices: continue 
                    if is_paragraph_needs_polishing(para):
                        paras_to_process.append((i, para.text))
                        
                total_paras = len(paras_to_process)
                if total_paras > 0:
                    _publish_message(channel_name, "block", f"📄 解析成功！共提取 {total_paras} 个核心段落待处理。\n🚀 【多线程+二次提纯引擎】已激活（初始阶段需后台思考，请稍候）...\n\n")
                    
                    def process_single_para(para_idx, text_content):
                        if len(text_content) > 600:
                            sub_chunks = split_text_into_chunks(text_content, max_chars=600)
                            polished_sub = []
                            for sc in sub_chunks:
                                polished_sub.append(sync_ai_request(text=sc, api_key=api_key, base_url=base_url, model_name=model_name, mode=task.mode))
                            return para_idx, "".join(polished_sub)
                        else:
                            return para_idx, sync_ai_request(text=text_content, api_key=api_key, base_url=base_url, model_name=model_name, mode=task.mode)

                    results_dict = {}
                    completed_count = 0
                    is_cancelled = False
                    
                    with ThreadPoolExecutor(max_workers=15) as executor:
                        future_to_para = {executor.submit(process_single_para, p[0], p[1]): p for p in paras_to_process}
                        for future in as_completed(future_to_para):
                            cancel_key = f"cancel:task:{task_id}"
                            if redis_client.exists(cancel_key):
                                is_cancelled = True
                                task.status = 'cancelled'
                                redis_client.delete(cancel_key)
                                break 
                            
                            idx, original_text = future_to_para[future]
                            try:
                                _, polished_text = future.result()
                                if not polished_text or len(polished_text.strip()) < 5: polished_text = original_text
                                results_dict[idx] = polished_text
                                completed_count += 1
                                redis_client.sadd(done_indices_key, idx)
                                _publish_message(channel_name, "progress", int((completed_count / total_paras) * 100))
                                _publish_message(channel_name, "block", f"⚡ [并发加速] 第 {completed_count}/{total_paras} 个段落已完成覆盖...\n")
                            except Exception:
                                results_dict[idx] = original_text 
                                
                    for idx, new_text in results_dict.items():
                        replace_paragraph_text(doc.paragraphs[idx], new_text)
                        
                    if is_cancelled:
                        temp_path = os.path.join('outputs', f"polished_temp_{task.id}.docx")
                        doc.save(temp_path)
                        task.result_file_path = temp_path
                        db.session.commit()
                        return
                
                output_path = os.path.join('outputs', f"polished_{task.id}.docx")
                doc.save(output_path)
                task.result_file_path = output_path
                redis_client.delete(done_indices_key) 
                redis_client.delete(f"docx_progress:task:{task_id}") 
                _publish_message(channel_name, "block", "\n\n🎉 全局高并发重排完毕，正在生成下载链接...\n")

            # ==========================================
            # 统一收尾操作
            # ==========================================
            task.status = 'completed'
            db.session.commit()
            
            is_long_task = (getattr(task, 'task_type', '') == 'docx') or (getattr(task, 'task_type', 'text') == 'text' and total_chunks > 1)
            if is_long_task:
                _publish_message(channel_name, "progress", 100) 
            
            if getattr(task, 'task_type', '') == 'docx':
                _publish_message(channel_name, "download", f"/api/tasks/download/{task.id}")
            
            _publish_message(channel_name, "done", "完成")
            print("🎉 [Worker] 任务彻底完成并退出！")
            
        except Exception as e:
            print(f"❌ [Worker] 发生严重报错退出: {str(e)}")
            db.session.rollback()
            task.status = 'failed'
            if getattr(task, 'task_type', 'text') == 'text' and 'full_text' in locals() and full_text: 
                task.polished_text = full_text
            elif getattr(task, 'task_type', '') == 'docx' and doc is not None:
                temp_path = os.path.join('outputs', f"polished_temp_{task.id}.docx")
                doc.save(temp_path)
                task.result_file_path = temp_path
            db.session.commit()
            _publish_error(channel_name, f"发生网络或解析异常: {str(e)}")

def _publish_message(channel_name, msg_type, content):
    payload = json.dumps({"type": msg_type, "content": content})
    redis_client.publish(channel_name, f"data: {payload}\n\n")

def _publish_error(channel_name, error_msg):
    payload = json.dumps({"type": "fatal", "content": error_msg})
    redis_client.publish(channel_name, f"data: {payload}\n\n")